from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from datetime import datetime

def export_draft_to_docx(draft_text: str, matter_name: str, sources: list) -> Document:
    """
    Export a draft document to DOCX format
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading(matter_name, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("")
    
    # Process draft text - handle sections
    sections = re.split(r'\*\*SECTION \d+:', draft_text)
    
    for section in sections:
        if not section.strip():
            continue
        
        # Extract section title if present
        lines = section.strip().split('\n')
        first_line = lines[0].strip()
        
        # Check if first line is a section header
        if first_line.endswith('**') or first_line.isupper():
            section_title = first_line.replace('**', '').strip()
            doc.add_heading(section_title, 1)
            content = '\n'.join(lines[1:])
        else:
            content = section
        
        # Add content paragraphs
        for para in content.split('\n\n'):
            if para.strip():
                # Handle bullet points
                if para.strip().startswith('-') or para.strip().startswith('•'):
                    for line in para.split('\n'):
                        if line.strip():
                            clean_line = line.strip().lstrip('-•').strip()
                            p = doc.add_paragraph(clean_line, style='List Bullet')
                else:
                    doc.add_paragraph(para.strip())
    
    # Add sources appendix
    if sources:
        doc.add_page_break()
        doc.add_heading('Source Documents', 1)
        
        for source in sources:
            doc.add_heading(f"[{source['citation']}] {source['document']}", 2)
            if source.get('page'):
                doc.add_paragraph(f"Page: {source['page']}")
            doc.add_paragraph(f"Relevance: {source['similarity']:.2%}")
            doc.add_paragraph(f"Excerpt: {source['snippet']}")
            doc.add_paragraph("")
    
    return doc

def export_research_to_markdown(research_output: str, matter_name: str, 
                                sources: list, query: str) -> str:
    """
    Export research report to Markdown format
    """
    md = f"# Research Report: {matter_name}\n\n"
    md += f"**Generated:** {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n\n"
    md += f"**Research Question:** {query}\n\n"
    md += "---\n\n"
    
    # Add main research content
    md += "## Research Findings\n\n"
    md += research_output + "\n\n"
    
    # Add sources appendix
    md += "---\n\n"
    md += "## Source Documents\n\n"
    
    for source in sources:
        md += f"### [{source['citation']}] {source['document']}\n\n"
        if source.get('page'):
            md += f"**Page:** {source['page']}\n\n"
        md += f"**Relevance Score:** {source['similarity']:.2%}\n\n"
        md += f"**Excerpt:**\n\n"
        md += f"> {source['snippet']}\n\n"
        md += "---\n\n"
    
    return md

def export_summary_to_markdown(summary_output: str, matter_name: str,
                               sources: list, mode: str) -> str:
    """
    Export summary to Markdown format
    """
    md = f"# {mode}: {matter_name}\n\n"
    md += f"**Generated:** {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n\n"
    md += "---\n\n"
    
    # Add summary content
    md += summary_output + "\n\n"
    
    # Add sources appendix
    md += "---\n\n"
    md += "## Supporting Sources\n\n"
    
    for source in sources:
        md += f"### [{source['citation']}] {source['document']}\n\n"
        if source.get('page'):
            md += f"**Page:** {source['page']}\n\n"
        md += f"**Relevance Score:** {source['similarity']:.2%}\n\n"
        md += f"**Excerpt:**\n\n"
        md += f"> {source['snippet']}\n\n"
        md += "---\n\n"
    
    return md

def create_matter_report(matter, documents: list, history: list) -> str:
    """
    Create a comprehensive matter report in Markdown
    """
    md = f"# Matter Report: {matter.matter_name}\n\n"
    md += f"**Client:** {matter.client_name}\n\n"
    md += f"**Created:** {matter.created_date[:10]}\n\n"
    md += f"**Last Modified:** {matter.last_modified[:10]}\n\n"
    md += f"**Description:** {matter.description}\n\n"
    md += "---\n\n"
    
    # Add documents section
    md += "## Associated Documents\n\n"
    if documents:
        for doc in documents:
            md += f"- **{doc.title}** ({doc.file_type})\n"
            if doc.page_count:
                md += f"  - Pages: {doc.page_count}\n"
            md += f"  - Uploaded: {doc.upload_date[:10]}\n\n"
    else:
        md += "No documents uploaded yet.\n\n"
    
    # Add history section
    md += "---\n\n"
    md += "## Matter History\n\n"
    
    if history:
        for entry in reversed(history):  # Most recent first
            md += f"### {entry.action_type.title()} - {entry.timestamp[:10]}\n\n"
            md += f"**Description:** {entry.description}\n\n"
            if entry.sources_used:
                md += f"**Sources Used:** {', '.join(f'[{s}]' for s in entry.sources_used)}\n\n"
            md += "---\n\n"
    else:
        md += "No actions recorded yet.\n\n"
    
    return md