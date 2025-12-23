import os
import uuid
import json
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import Optional
import PyPDF2
import docx
from vector_store import VectorStore

@dataclass
class Document:
    doc_id: str
    title: str
    file_type: str
    text: str
    page_count: Optional[int] = None
    upload_date: str = None
    original_filename: str = None
    
    def to_dict(self):
        return asdict(self)

class DocumentProcessor:
    def __init__(self, storage_dir="documents"):
        self.storage_dir = Path(storage_dir)
        self.raw_dir = self.storage_dir / "raw"
        self.extracted_dir = self.storage_dir / "extracted"
        
        # Create directories if they don't exist
        self.raw_dir.mkdir(parents=True, exist_ok=True)
        self.extracted_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize vector store
        self.vector_store = VectorStore()
    
    def process_file(self, uploaded_file) -> Document:
        """Process uploaded file and return Document object"""
        doc_id = str(uuid.uuid4())
        original_filename = uploaded_file.name
        file_extension = Path(original_filename).suffix.lower()
        
        # Save raw file
        raw_path = self.raw_dir / f"{doc_id}{file_extension}"
        with open(raw_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text, page_count = self._extract_pdf(raw_path)
            file_type = "PDF"
        elif file_extension == '.docx':
            text = self._extract_docx(raw_path)
            page_count = None
            file_type = "DOCX"
        elif file_extension == '.txt':
            text = self._extract_txt(raw_path)
            page_count = None
            file_type = "TXT"
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Create document object
        doc = Document(
            doc_id=doc_id,
            title=Path(original_filename).stem,
            file_type=file_type,
            text=text,
            page_count=page_count,
            upload_date=datetime.now().isoformat(),
            original_filename=original_filename
        )
        
        # Save extracted text and metadata
        self._save_extracted(doc)
        
        # Add to vector store
        num_chunks = self.vector_store.add_document(
            doc_id=doc.doc_id,
            doc_title=doc.title,
            text=doc.text,
            page_count=doc.page_count
        )
        print(f"Added {num_chunks} chunks to vector store")
        
        return doc
    
    def _extract_pdf(self, file_path) -> tuple[str, int]:
        """Extract text from PDF and return text + page count"""
        text_parts = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                # Clean up common PDF extraction issues
                page_text = self._clean_extracted_text(page_text)
                text_parts.append(f"--- Page {page_num} ---\n{page_text}\n")
        
        return "\n".join(text_parts), page_count
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean up common PDF text extraction issues"""
        import re
        
        # Fix numbers that are mashed together (e.g., "8.99are" -> "8.99 are")
        text = re.sub(r'(\d+\.?\d*)([a-zA-Z])', r'\1 \2', text)
        
        # Fix words mashed together with numbers (e.g., "below8.99" -> "below 8.99")
        text = re.sub(r'([a-zA-Z])(\d+\.?\d*)', r'\1 $\2', text)
        
        # Fix multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Fix common ligatures that get extracted incorrectly
        text = text.replace('ﬁ', 'fi')
        text = text.replace('ﬂ', 'fl')
        text = text.replace('ﬀ', 'ff')
        text = text.replace('ﬃ', 'ffi')
        text = text.replace('ﬄ', 'ffl')
        
        return text
    
    def _extract_docx(self, file_path) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        return "\n\n".join(text_parts)
    
    def _extract_txt(self, file_path) -> str:
        """Extract text from TXT"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _save_extracted(self, doc: Document):
        """Save extracted text and metadata to disk"""
        # Save text file
        text_path = self.extracted_dir / f"{doc.doc_id}.txt"
        with open(text_path, 'w', encoding='utf-8') as f:
            f.write(doc.text)
        
        # Save metadata as JSON
        metadata_path = self.extracted_dir / f"{doc.doc_id}_metadata.json"
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(doc.to_dict(), f, indent=2)
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """Retrieve a document by ID"""
        metadata_path = self.extracted_dir / f"{doc_id}_metadata.json"
        
        if not metadata_path.exists():
            return None
        
        with open(metadata_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return Document(**data)
    
    def list_documents(self) -> list[Document]:
        """List all processed documents"""
        documents = []
        for metadata_file in self.extracted_dir.glob("*_metadata.json"):
            with open(metadata_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                documents.append(Document(**data))
        
        # Sort by upload date (newest first)
        documents.sort(key=lambda x: x.upload_date, reverse=True)
        return documents
    
    def search_documents(self, query: str, top_k: int = 5):
        """
        Search for top-k most relevant chunks across all documents
        Returns list of results with doc_title, chunk_text, page, and similarity
        """
        return self.vector_store.search(query, top_k=top_k)
    
    def get_vector_stats(self):
        """Get statistics about the vector store"""
        return self.vector_store.get_stats()
    
    def delete_document(self, doc_id: str) -> bool:
        """Delete a document and all its associated files"""
        try:
            # Find the original file extension from metadata
            metadata_path = self.extracted_dir / f"{doc_id}_metadata.json"
            
            if not metadata_path.exists():
                return False
            
            # Load metadata to get original filename and extension
            with open(metadata_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                original_filename = data.get('original_filename', '')
                file_extension = Path(original_filename).suffix.lower()
            
            # Delete raw file
            raw_path = self.raw_dir / f"{doc_id}{file_extension}"
            if raw_path.exists():
                os.remove(raw_path)
            
            # Delete extracted text file
            text_path = self.extracted_dir / f"{doc_id}.txt"
            if text_path.exists():
                os.remove(text_path)
            
            # Delete metadata file
            if metadata_path.exists():
                os.remove(metadata_path)
            
            # Delete from vector store
            self.vector_store.delete_document(doc_id)
            
            return True
        except Exception as e:
            print(f"Error deleting document {doc_id}: {str(e)}")
            return False